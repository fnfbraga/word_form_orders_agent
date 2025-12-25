#!/usr/bin/env python3
"""Script to create a sample movie order form template."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_order_form_template():
    """Create a sample movie order form template."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Movie Order Form', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Please fill out this form to order your movies')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Customer Information Section
    doc.add_heading('Customer Information', level=1)
    
    # Add form fields
    doc.add_paragraph('Name: ________________________________________')
    doc.add_paragraph('Street: ________________________________________')
    doc.add_paragraph('Postal Code and City: ________________________________________')
    doc.add_paragraph('Country: ________________________________________')
    
    doc.add_paragraph()  # Spacer
    
    # Movies Section
    doc.add_heading('Movies Ordered', level=1)
    
    # Create table for movies
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = Inches(3.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(2)
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Title'
    header_cells[1].text = 'Language'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()  # Spacer
    
    # Footer
    footer = doc.add_paragraph('Thank you for your order!')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    doc.save('OrderForm.docx')
    print('Created OrderForm.docx')


if __name__ == '__main__':
    create_order_form_template()


