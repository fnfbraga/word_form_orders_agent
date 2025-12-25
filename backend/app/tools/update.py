"""Tool for updating the order document with collected data."""

import re
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.table import Table
from docx.oxml.ns import qn
from lxml import etree
try:
    from mailmerge import MailMerge
    HAS_MAILMERGE = True
except ImportError:
    HAS_MAILMERGE = False
    MailMerge = None

from app.session import Session


# Checkbox unicode characters
UNCHECKED_BOXES = ['☐', '□', '○', '◯', '( )', '[ ]']
CHECKED_BOX = '☑'


def update_order_document(
    session: Session,
    name: str | None = None,
    street: str | None = None,
    postal_code_city: str | None = None,
    country: str | None = None,
    movies: list[dict[str, str]] | None = None
) -> dict:
    """
    Update the session's form data with provided values.
    
    Args:
        session: The user session
        name: Customer name
        street: Street address
        postal_code_city: Postal code and city
        country: Country (should be pre-validated)
        movies: List of movies with title and language
        
    Returns:
        Dictionary with update status and current form state
    """
    form_data = session.form_data
    
    # Update provided fields
    if name is not None:
        form_data.name = name
    if street is not None:
        form_data.street = street
    if postal_code_city is not None:
        form_data.postal_code_city = postal_code_city
    if country is not None:
        form_data.country = country
    if movies is not None:
        form_data.movies = movies
    
    # Check completion status
    session.is_complete = form_data.is_complete()
    
    # Build response
    missing_fields = []
    if not form_data.name:
        missing_fields.append("Name")
    if not form_data.street:
        missing_fields.append("Street")
    if not form_data.postal_code_city:
        missing_fields.append("Postal Code and City")
    if not form_data.country:
        missing_fields.append("Country")
    if not form_data.movies:
        missing_fields.append("Movies")
    
    return {
        "success": True,
        "is_complete": session.is_complete,
        "current_data": form_data.to_dict(),
        "missing_fields": missing_fields,
        "message": "Form is complete!" if session.is_complete else f"Still need: {', '.join(missing_fields)}"
    }


def add_movie(session: Session, title: str, language: str) -> dict:
    """
    Add a movie to the order.
    
    Args:
        session: The user session
        title: Movie title
        language: Movie language
        
    Returns:
        Dictionary with current movie list
    """
    movie = {"title": title.strip(), "language": language.strip()}
    session.form_data.movies.append(movie)
    
    # Check completion status
    session.is_complete = session.form_data.is_complete()
    
    return {
        "success": True,
        "movie_added": movie,
        "total_movies": len(session.form_data.movies),
        "movies": session.form_data.movies,
        "message": f"Added '{title}' ({language}). Total: {len(session.form_data.movies)} movie(s)."
    }


def _get_merge_fields(file_path: Path) -> list[str]:
    """Get list of merge field names in the document."""
    if not HAS_MAILMERGE:
        return []
    try:
        with MailMerge(file_path) as doc:
            return list(doc.get_merge_fields())
    except Exception:
        return []


def _fill_with_mailmerge(file_path: Path, output_path: Path, form_data, movies: list[dict]) -> bool:
    """
    Try to fill document using mail merge fields.
    
    Returns True if successful, False if no merge fields found.
    """
    if not HAS_MAILMERGE:
        return False
    try:
        with MailMerge(file_path) as doc:
            merge_fields = doc.get_merge_fields()
            
            if not merge_fields:
                return False
            
            # Build merge data - map common field names
            merge_data = {}
            
            # Customer fields - try various common names
            name_fields = ['Name', 'name', 'CustomerName', 'customer_name', 'FullName', 'full_name']
            street_fields = ['Street', 'street', 'Address', 'address', 'StreetAddress', 'street_address']
            postal_fields = ['PostalCodeCity', 'postal_code_city', 'PostalCode', 'postal_code', 
                           'City', 'city', 'ZipCity', 'zip_city', 'Zip', 'zip']
            country_fields = ['Country', 'country', 'Nation', 'nation']
            
            for field in merge_fields:
                if field in name_fields and form_data.name:
                    merge_data[field] = form_data.name
                elif field in street_fields and form_data.street:
                    merge_data[field] = form_data.street
                elif field in postal_fields and form_data.postal_code_city:
                    merge_data[field] = form_data.postal_code_city
                elif field in country_fields and form_data.country:
                    merge_data[field] = form_data.country
            
            # Movie fields - handle multiple movies
            for i, movie in enumerate(movies):
                idx = i + 1
                title = movie.get("title", "")
                language = movie.get("language", "")
                
                # Try various movie field naming conventions
                title_names = [f'MovieTitle{idx}', f'movie_title_{idx}', f'Title{idx}', f'title_{idx}',
                              'MovieTitle', 'movie_title', 'Title', 'title', 'MovieName', 'movie_name']
                lang_names = [f'MovieLanguage{idx}', f'movie_language_{idx}', f'Language{idx}', f'language_{idx}',
                             'MovieLanguage', 'movie_language', 'Language', 'language', 'Lang', 'lang']
                
                for field in merge_fields:
                    if field in title_names and title:
                        merge_data[field] = title
                    elif field in lang_names and language:
                        merge_data[field] = language
            
            if merge_data:
                doc.merge(**merge_data)
                doc.write(output_path)
                return True
            
            return False
    except Exception as e:
        print(f"Mailmerge error: {e}")
        return False


def _fill_content_controls(doc: Document, form_data, movies: list[dict]) -> bool:
    """
    Fill content control (SDT) fields in the document.
    
    Returns True if any fields were filled.
    """
    filled = False
    
    # XML namespaces
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
    }
    
    # Field mappings (tag/alias -> value)
    field_values = {}
    if form_data.name:
        field_values.update({'name': form_data.name, 'Name': form_data.name, 
                            'CustomerName': form_data.name, 'customer_name': form_data.name})
    if form_data.street:
        field_values.update({'street': form_data.street, 'Street': form_data.street,
                            'address': form_data.street, 'Address': form_data.street})
    if form_data.postal_code_city:
        field_values.update({'postal_code_city': form_data.postal_code_city, 
                            'PostalCodeCity': form_data.postal_code_city,
                            'city': form_data.postal_code_city, 'City': form_data.postal_code_city,
                            'postal_code': form_data.postal_code_city, 'PostalCode': form_data.postal_code_city,
                            'zip': form_data.postal_code_city, 'Zip': form_data.postal_code_city})
    if form_data.country:
        field_values.update({'country': form_data.country, 'Country': form_data.country})
    
    # Add movie fields
    if movies:
        movie = movies[0]  # First movie for single-movie fields
        title = movie.get("title", "")
        language = movie.get("language", "")
        if title:
            field_values.update({'MovieTitle': title, 'movie_title': title, 
                                'Title': title, 'title': title,
                                'MovieName': title, 'movie_name': title})
        if language:
            field_values.update({'MovieLanguage': language, 'movie_language': language,
                                'Language': language, 'language': language,
                                'Lang': language, 'lang': language})
    
    # Find and fill SDT elements
    for sdt in doc.element.iter(qn('w:sdt')):
        # Get the tag/alias of this content control
        sdt_pr = sdt.find('.//w:sdtPr', nsmap)
        if sdt_pr is None:
            continue
        
        # Check tag element
        tag_elem = sdt_pr.find('.//w:tag', nsmap)
        alias_elem = sdt_pr.find('.//w:alias', nsmap)
        
        tag_name = None
        if tag_elem is not None:
            tag_name = tag_elem.get(qn('w:val'))
        elif alias_elem is not None:
            tag_name = alias_elem.get(qn('w:val'))
        
        if tag_name and tag_name in field_values:
            # Find the text content element and update it
            sdt_content = sdt.find('.//w:sdtContent', nsmap)
            if sdt_content is not None:
                # Find text runs and update
                for run in sdt_content.iter(qn('w:r')):
                    for text_elem in run.iter(qn('w:t')):
                        text_elem.text = field_values[tag_name]
                        filled = True
                        break
    
    return filled


def _replace_unicode_checkbox(text: str, should_check: bool = True) -> str:
    """Replace unicode unchecked box with checked box."""
    result = text
    for unchecked in UNCHECKED_BOXES:
        if unchecked in result:
            result = result.replace(unchecked, CHECKED_BOX if should_check else unchecked)
    return result


def _tick_movie_checkboxes(doc: Document, movie_titles: list[str]) -> None:
    """Find and tick checkboxes next to movie titles in the document."""
    titles_lower = [t.lower().strip() for t in movie_titles]
    
    # Check paragraphs for checkbox patterns
    for para in doc.paragraphs:
        text = para.text
        text_lower = text.lower()
        
        for title in titles_lower:
            if title in text_lower:
                for run in para.runs:
                    run.text = _replace_unicode_checkbox(run.text)
                break
    
    # Check tables for checkbox patterns
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text for cell in row.cells).lower()
            for title in titles_lower:
                if title in row_text:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = _replace_unicode_checkbox(run.text)
                    break


def _fill_text_placeholders(doc: Document, form_data, movies: list[dict]) -> None:
    """
    Fill document by replacing text placeholders.
    Section-aware to avoid mixing customer and movie fields.
    """
    current_section = "customer"
    movie_idx = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detect section changes
        if re.search(r"(?i)movies?\s*ordered", text):
            current_section = "movies"
            continue
        elif re.search(r"(?i)customer\s*info", text):
            current_section = "customer"
            continue
        
        original_text = para.text
        new_text = original_text
        
        if current_section == "customer":
            # Replace customer fields
            if form_data.name and re.search(r"(?i)^name\s*[:：]", text):
                new_text = re.sub(r"(?i)(name\s*[:：]\s*)(.*)", r"\1" + form_data.name, new_text)
            
            if form_data.street and re.search(r"(?i)^street\s*[:：]", text):
                new_text = re.sub(r"(?i)(street\s*[:：]\s*)(.*)", r"\1" + form_data.street, new_text)
            
            if form_data.postal_code_city:
                if re.search(r"(?i)^(postal|zip|city)", text):
                    new_text = re.sub(
                        r"(?i)((?:postal\s*code|zip\s*code?|city)(?:\s*(?:and|&)\s*(?:city|postal\s*code))?\s*[:：]\s*)(.*)",
                        r"\1" + form_data.postal_code_city,
                        new_text
                    )
            
            if form_data.country and re.search(r"(?i)^country\s*[:：]", text):
                new_text = re.sub(r"(?i)(country\s*[:：]\s*)(.*)", r"\1" + form_data.country, new_text)
        
        elif current_section == "movies" and movies and movie_idx < len(movies):
            movie = movies[movie_idx]
            
            # Handle combined Name: ... Language: ... on same line
            if re.search(r"(?i)(name|title)\s*[:：]", text) and re.search(r"(?i)language\s*[:：]", text):
                title = movie.get("title", "")
                language = movie.get("language", "")
                # Replace both on the same line
                new_text = re.sub(
                    r"(?i)((name|title)\s*[:：]\s*)([^:]*?)(language\s*[:：]\s*)(.*)",
                    r"\1" + title + r"\4" + language,
                    new_text
                )
                movie_idx += 1
            else:
                # Handle separate lines
                if re.search(r"(?i)^(name|title)\s*[:：]", text):
                    title = movie.get("title", "")
                    new_text = re.sub(r"(?i)((name|title)\s*[:：]\s*)(.*)", r"\1" + title, new_text)
                
                if re.search(r"(?i)^language\s*[:：]", text):
                    language = movie.get("language", "")
                    new_text = re.sub(r"(?i)(language\s*[:：]\s*)(.*)", r"\1" + language, new_text)
                    movie_idx += 1
        
        # Apply changes
        if new_text != original_text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)


def _fill_movie_table(table: Table, movies: list[dict[str, str]]) -> None:
    """Fill a structured movie table with Title/Language columns."""
    title_col = None
    language_col = None
    name_col = None
    
    if len(table.rows) > 0:
        for idx, cell in enumerate(table.rows[0].cells):
            header = cell.text.strip().lower()
            if "title" in header:
                title_col = idx
            elif "name" in header and title_col is None:
                name_col = idx
            elif "language" in header or "lang" in header:
                language_col = idx
    
    # Use name column if no title column
    if title_col is None:
        title_col = name_col
    
    if title_col is None or language_col is None:
        return
    
    # Remove existing data rows (keep header)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)
    
    # Add movies
    for movie in movies:
        row = table.add_row()
        row.cells[title_col].text = movie.get("title", "")
        row.cells[language_col].text = movie.get("language", "")


def generate_filled_document(session: Session) -> Path:
    """
    Generate the filled document with all form data.
    
    Uses multiple strategies:
    1. Mail merge fields (if present)
    2. Content controls / form fields
    3. Text placeholder replacement
    """
    form_data = session.form_data
    structure = session.form_structure
    movies = form_data.movies or []
    
    output_path = session.file_path.parent / f"filled_{session.file_path.name}"
    
    # Strategy 1: Try mail merge first
    if _fill_with_mailmerge(session.file_path, output_path, form_data, movies):
        # Mailmerge worked, but we may need to post-process for checkboxes
        doc = Document(output_path)
        if movies:
            movie_titles = [m.get("title", "") for m in movies]
            _tick_movie_checkboxes(doc, movie_titles)
        doc.save(output_path)
        return output_path
    
    # Strategy 2 & 3: Use python-docx for content controls and text replacement
    doc = Document(session.file_path)
    
    # Try content controls first
    _fill_content_controls(doc, form_data, movies)
    
    # Always do text placeholder replacement as fallback
    _fill_text_placeholders(doc, form_data, movies)
    
    # Tick checkboxes for ordered movies
    if movies:
        movie_titles = [m.get("title", "") for m in movies]
        _tick_movie_checkboxes(doc, movie_titles)
    
    # Handle structured movie table
    if structure and structure.has_movie_table and structure.movie_table_index is not None:
        table = doc.tables[structure.movie_table_index]
        _fill_movie_table(table, movies)
    
    doc.save(output_path)
    return output_path

