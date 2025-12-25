"""Tool for inspecting form structure in uploaded documents."""

import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from app.session import Session, FormStructure


# Checkbox unicode characters to detect
CHECKBOX_CHARS = ['☐', '□', '○', '◯', '☑', '✓', '✔']


def _extract_movie_options(doc: Document) -> list[str]:
    """
    Extract available movie options from checkbox lists in the document.
    
    Args:
        doc: The Word document
        
    Returns:
        List of movie titles found next to checkboxes
    """
    movies = []
    checkbox_pattern = r'[☐□○◯☑✓✔\[\]\(\)]'
    
    # Check paragraphs for checkbox + movie patterns
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Look for checkbox characters followed by text
        if re.search(checkbox_pattern, text):
            # Extract the text after the checkbox
            movie_text = re.sub(checkbox_pattern, '', text).strip()
            # Clean up common separators
            movie_text = re.sub(r'^[\s\-–—:]+', '', movie_text).strip()
            if movie_text and len(movie_text) > 1:
                movies.append(movie_text)
    
    # Check tables for checkbox patterns
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text.strip() for cell in row.cells)
            if re.search(checkbox_pattern, row_text):
                # Extract movie name from the row
                movie_text = re.sub(checkbox_pattern, '', row_text).strip()
                movie_text = re.sub(r'^[\s\-–—:]+', '', movie_text).strip()
                if movie_text and len(movie_text) > 1:
                    movies.append(movie_text)
    
    return movies


def inspect_form_structure(session: Session) -> dict:
    """
    Inspect the uploaded document to identify form fields and structure.
    
    Returns a dictionary describing the form structure including:
    - Detected placeholder fields (Name, Street, etc.)
    - Whether a movie table exists
    - Whether a checkbox movie list exists
    - Available movie options if present
    """
    doc = Document(session.file_path)
    structure = FormStructure()
    
    # Patterns to match common field labels
    field_patterns = {
        "name": r"(?i)\b(name|full\s*name|customer\s*name)\s*[:：]?\s*[_\s]*",
        "street": r"(?i)\b(street|address|street\s*address)\s*[:：]?\s*[_\s]*",
        "postal_city": r"(?i)\b(postal\s*code|zip|city|postal\s*code\s*and\s*city|zip\s*code)\s*[:：]?\s*[_\s]*",
        "country": r"(?i)\b(country|nation)\s*[:：]?\s*[_\s]*"
    }
    
    # Scan paragraphs for field placeholders
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        for field_name, pattern in field_patterns.items():
            if re.search(pattern, text):
                if field_name == "name":
                    structure.has_name_field = True
                    structure.placeholders["name"] = text
                elif field_name == "street":
                    structure.has_street_field = True
                    structure.placeholders["street"] = text
                elif field_name == "postal_city":
                    structure.has_postal_city_field = True
                    structure.placeholders["postal_city"] = text
                elif field_name == "country":
                    structure.has_country_field = True
                    structure.placeholders["country"] = text
    
    # Scan tables for movie list
    for table_idx, table in enumerate(doc.tables):
        headers = []
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip().lower())
        
        # Check if this looks like a movie table
        has_title = any("title" in h for h in headers)
        has_language = any("language" in h or "lang" in h for h in headers)
        
        if has_title and has_language:
            structure.has_movie_table = True
            structure.movie_table_index = table_idx
            break
    
    # Check for checkbox-based movie list
    movie_options = _extract_movie_options(doc)
    has_checkbox_list = len(movie_options) > 0
    
    # Store checkbox info in structure
    structure.has_checkbox_list = has_checkbox_list
    structure.available_movies = movie_options
    
    # Store structure in session
    session.form_structure = structure
    
    # Build response for the agent
    detected_fields = []
    if structure.has_name_field:
        detected_fields.append("Name")
    if structure.has_street_field:
        detected_fields.append("Street/Address")
    if structure.has_postal_city_field:
        detected_fields.append("Postal Code/City")
    if structure.has_country_field:
        detected_fields.append("Country")
    
    # Build movie section message
    movie_info = []
    if structure.has_movie_table:
        movie_info.append("a movie table for adding custom movies")
    if has_checkbox_list:
        movie_info.append(f"a checkbox list with {len(movie_options)} movie options: {', '.join(movie_options[:10])}")
    
    movie_message = ' and '.join(movie_info) if movie_info else 'no movie section'
    
    return {
        "success": True,
        "detected_fields": detected_fields,
        "has_movie_table": structure.has_movie_table,
        "has_checkbox_list": has_checkbox_list,
        "available_movies": movie_options if has_checkbox_list else [],
        "message": f"Found {len(detected_fields)} address fields and {movie_message}."
    }

